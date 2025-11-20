"""
Document loading service for multiple file formats.
"""
from pathlib import Path
from typing import List
import PyPDF2
from docx import Document as DocxDocument
from PIL import Image
import pytesseract
from langchain.schema import Document
from utils.logger import logger
from utils.config import settings


class DocumentLoader:
    """Load documents from supported formats."""

    def __init__(self) -> None:
        self.supported_formats = settings.allowed_extensions
        logger.info("Supported formats: %s", self.supported_formats)

    def load_pdf(self, file_path: str) -> tuple[str, int]:
        """Extract text from PDF."""
        text = ""
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            page_count = len(pdf_reader.pages)

            for page_num, page in enumerate(pdf_reader.pages):
                text += f"\n--- Page {page_num + 1} ---\n"
                extracted = page.extract_text() or ""
                text += extracted

        logger.info("Extracted %s pages from PDF %s", page_count, file_path)
        return text, page_count

    def load_docx(self, file_path: str) -> tuple[str, int]:
        """Extract text from DOCX."""
        doc = DocxDocument(file_path)
        text = "\n".join(para.text for para in doc.paragraphs)
        paragraph_count = len(doc.paragraphs)
        logger.info("Extracted %s paragraphs from DOCX %s", paragraph_count, file_path)
        return text, paragraph_count

    def load_text(self, file_path: str) -> tuple[str, int]:
        """Load plain text file."""
        with open(file_path, "r", encoding="utf-8") as file:
            text = file.read()
        line_count = len(text.splitlines())
        logger.info("Loaded text file %s (%s lines)", file_path, line_count)
        return text, line_count

    def extract_text_from_image(self, file_path: str) -> str:
        """Extract text from an image via OCR."""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            logger.info("Extracted text from image %s", file_path)
            return text
        except pytesseract.TesseractNotFoundError:
            error_msg = (
                "Tesseract OCR is not installed or not in your PATH. "
                "Image processing requires Tesseract OCR. "
                "See backend/TESSERACT_INSTALL.md for installation instructions. "
                "PDF, DOCX, and TXT files work without Tesseract."
            )
            logger.error(error_msg)
            raise RuntimeError(error_msg)
        except Exception as exc:
            logger.error("Failed to extract text from image %s: %s", file_path, exc)
            raise

    def load_document(self, file_path: str):
        """Load a document and return a LangChain Document."""
        path = Path(file_path)
        file_extension = path.suffix.lower().lstrip(".")
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")

        file_name = path.name

        if file_extension == "pdf":
            content, pages = self.load_pdf(file_path)
            metadata = {"source": file_name, "file_type": "pdf", "pages": pages}
        elif file_extension == "docx":
            content, paragraphs = self.load_docx(file_path)
            metadata = {
                "source": file_name,
                "file_type": "docx",
                "paragraphs": paragraphs,
            }
        elif file_extension == "txt":
            content, lines = self.load_text(file_path)
            metadata = {"source": file_name, "file_type": "txt", "lines": lines}
        elif file_extension in ["png", "jpg", "jpeg"]:
            content = self.extract_text_from_image(file_path)
            metadata = {
                "source": file_name,
                "file_type": "image",
                "format": file_extension,
            }
        else:
            raise ValueError(f"Unsupported format: {file_extension}")

        return Document(page_content=content, metadata=metadata)

    def load_documents_from_directory(self, directory: str) -> List[Document]:
        """Load all supported documents from a directory."""
        documents: List[Document] = []
        dir_path = Path(directory)
        if not dir_path.exists():
            logger.warning("Directory not found: %s", directory)
            return documents

        for file_path in dir_path.glob("*"):
            if file_path.is_file() and file_path.suffix.lower().lstrip(
                "."
            ) in self.supported_formats:
                try:
                    document = self.load_document(str(file_path))
                    documents.append(document)
                    logger.info("Loaded document: %s", file_path.name)
                except Exception as exc:
                    logger.error("Failed to load %s: %s", file_path.name, exc)
        logger.info("Loaded %s documents from %s", len(documents), directory)
        return documents

